# -*- coding: utf-8 -*-
"""
Generates the research project report (.docx) for the Online Mental Health
Self-Assessment System (case study in Tanzania), formatted to follow the
Kampala International University in Tanzania (KIUT) Department of Computing and
Information Technology project guidelines:

  * Times New Roman; body 12pt; topics/titles 14pt; bold (not underlined);
    1.5 line spacing.
  * Margins: 1" top/bottom/right; 1.25" left (for binding).
  * Pagination centered at the bottom; preliminary pages in lower-case Roman
    numerals (starting at ii on the Declaration); main document in Arabic
    numerals restarting at 1 on Chapter One; the title page is not paginated.
  * Preliminary order: Title, Declaration, Dedication, Acknowledgement,
    Abstract, Table of Contents, List of Tables, List of Figures, List of
    Acronyms/Abbreviations, Definition of Key Terms.
  * Chapters One to Five following the prescribed section outline, References
    (APA 7th), and Appendices.

The content reflects the system as actually implemented: MindCare branding with
JWT user authentication (register/login), an administrator panel for user
management and report generation, downloadable mental-health reports (.docx),
three validated instruments (PHQ-9, GAD-7, SDQ), AI/LLM-generated feedback and
the "Mia" assistant, bilingual English/Swahili, on a Vue 3 + Django REST stack
with SQLite (development) / PostgreSQL (production).
"""

import os

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


DIAGRAMS = r"d:\Careen\frontend\diagrams"

doc = Document()

# ----------------------------------------------------------------------------- base styles
normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"
normal.font.size = Pt(12)
normal.paragraph_format.space_after = Pt(8)
normal.paragraph_format.line_spacing = 1.5

BLACK = RGBColor(0, 0, 0)
for name, size, centered in (
    ("Heading 1", 14, True),
    ("Heading 2", 14, False),
    ("Heading 3", 13, False),
    ("Title", 14, True),
):
    try:
        st = doc.styles[name]
        st.font.name = "Times New Roman"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.italic = False
        st.font.underline = False
        st.font.color.rgb = BLACK
        st.paragraph_format.space_before = Pt(10)
        st.paragraph_format.space_after = Pt(6)
        if centered:
            st.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            st.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    except KeyError:
        pass

# base margins on the first section (title page)
sec0 = doc.sections[0]
for s in (sec0,):
    s.top_margin = Inches(1)
    s.bottom_margin = Inches(1)
    s.right_margin = Inches(1)
    s.left_margin = Inches(1.25)


# ----------------------------------------------------------------------------- helpers
def apply_margins(section):
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.right_margin = Inches(1)
    section.left_margin = Inches(1.25)


def set_page_number_format(section, fmt, start):
    sectPr = section._sectPr
    pgNumType = sectPr.find(qn("w:pgNumType"))
    if pgNumType is None:
        pgNumType = OxmlElement("w:pgNumType")
        sectPr.append(pgNumType)
    pgNumType.set(qn("w:fmt"), fmt)
    pgNumType.set(qn("w:start"), str(start))


def add_footer_page_number(section):
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    b = OxmlElement("w:fldChar"); b.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = "PAGE"
    e = OxmlElement("w:fldChar"); e.set(qn("w:fldCharType"), "end")
    run._r.append(b); run._r.append(instr); run._r.append(e)


def h1(text):
    return doc.add_heading(text, level=1)


def h2(text):
    return doc.add_heading(text, level=2)


def h3(text):
    return doc.add_heading(text, level=3)


def para(text="", bold=False, italic=False, align="justify", size=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    p.alignment = {
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }[align]
    return p


def roman_item(numeral, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    r = p.add_run(numeral + ".  ")
    r.bold = True
    p.add_run(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def bullet(text):
    return doc.add_paragraph(text, style="List Bullet")


def numbered(text):
    return doc.add_paragraph(text, style="List Number")


def caption(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    return p


def figure_caption(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    return p


def placeholder(note):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("[ " + note + " ]")
    run.italic = True
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    return p


def add_diagram(filename, width_inches=6.2):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(os.path.join(DIAGRAMS, filename), width=Inches(width_inches))
    return p


def add_screenshot(filename, width_inches=5.5):
    """Embed a captured application screenshot from the screens sub-folder."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(os.path.join(DIAGRAMS, "screens", filename), width=Inches(width_inches))
    return p


def code_listing(code, caption_text=None):
    """Render a source-code excerpt in a bordered, single-cell, monospaced block."""
    if caption_text:
        c = doc.add_paragraph()
        cr = c.add_run(caption_text)
        cr.bold = True
        cr.font.size = Pt(11)
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    cell.paragraphs[0].text = ""
    first = True
    for line in code.split("\n"):
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(line if line else " ")
        r.font.name = "Consolas"
        r.font.size = Pt(8.5)
    doc.add_paragraph()
    return table


def page_break():
    doc.add_page_break()


def add_toc():
    p = doc.add_paragraph()
    run = p.add_run()
    b = OxmlElement("w:fldChar"); b.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    sep = OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t"); t.text = "Right-click and choose \u201cUpdate Field\u201d to generate the Table of Contents."
    e = OxmlElement("w:fldChar"); e.set(qn("w:fldCharType"), "end")
    for x in (b, instr, sep, t, e):
        run._r.append(x)


def simple_table(headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ""
        r = hdr[i].paragraphs[0].add_run(htext)
        r.bold = True
        r.font.size = Pt(11)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            rr = cells[i].paragraphs[0].add_run(str(val))
            rr.font.size = Pt(10)
    doc.add_paragraph()
    return table


# ============================================================================= TITLE PAGE
for _ in range(2):
    doc.add_paragraph()
para("ONLINE MENTAL HEALTH SELF-ASSESSMENT SYSTEM", bold=True, align="center", size=14)
para("A CASE STUDY IN TANZANIA", bold=True, align="center", size=14)
doc.add_paragraph()
doc.add_paragraph()
para("CAREEN GODSON BIRDI", bold=True, align="center", size=14)
para("BCS/30861/2301/DT", bold=True, align="center", size=14)
doc.add_paragraph()
doc.add_paragraph()
para(
    "Research Project Submitted in Partial Fulfillment for the Award of the "
    "Degree of Bachelor of Computer Science at Kampala International University "
    "in Tanzania",
    align="center",
)
doc.add_paragraph()
doc.add_paragraph()
para("June, 2026", bold=True, align="center")

# ----------------------------------------------- section break -> preliminary (roman ii)
sec_prelim = doc.add_section(WD_SECTION.NEW_PAGE)
apply_margins(sec_prelim)
set_page_number_format(sec_prelim, "lowerRoman", 2)
add_footer_page_number(sec_prelim)

# ============================================================================= DECLARATION
h1("DECLARATION")
para(
    "This project is my original work and has not been presented for a degree in "
    "any other university or for any other award."
)
doc.add_paragraph()
para("Student Name: Careen Godson Birdi")
para("Sign: \u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026   Date: \u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026")
doc.add_paragraph()
para(
    "I confirm that the work reported in this project was carried out by the "
    "candidate under my supervision."
)
doc.add_paragraph()
para("Supervisor Name: \u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026")
para("Sign: \u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026   Date: \u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026")
page_break()

# ============================================================================= DEDICATION
h1("DEDICATION")
para(
    "To my beloved parents and family, for their endless support, encouragement, "
    "and motivation throughout my academic journey."
)
page_break()

# ============================================================================= ACKNOWLEDGEMENT
h1("ACKNOWLEDGEMENT")
para(
    "I sincerely thank Almighty God for the strength and knowledge to complete this "
    "project. I am deeply grateful to my project supervisor for the valuable "
    "guidance, constructive feedback, and continuous support provided throughout the "
    "research. I also thank the lecturers of the Department of Computing and "
    "Information Technology at Kampala International University in Tanzania for the "
    "knowledge and skills imparted during my studies. My appreciation goes to the "
    "mental-health professionals, ICT specialists, and university students who "
    "participated in this study and provided insights that shaped the development of "
    "the system. Finally, I thank my family and friends for their encouragement, "
    "patience, and moral support throughout this journey. May God bless everyone who "
    "contributed, directly or indirectly, to the successful completion of this work."
)
page_break()

# ============================================================================= ABSTRACT
h1("ABSTRACT")
abs_p = doc.add_paragraph()
abs_p.paragraph_format.line_spacing = 1.0
abs_run = abs_p.add_run(
    "The purpose of this study was to design and implement an online mental-health "
    "self-assessment system that supports early detection of mental-health concerns "
    "and delivers immediate, artificial-intelligence-assisted feedback and referral "
    "guidance for users in Tanzania. The specific objectives were to analyse user "
    "requirements, design the architecture and bilingual interfaces, develop a "
    "web-based system offering validated assessments (PHQ-9, GAD-7, and SDQ) with "
    "AI-generated feedback and downloadable reports, integrate a conversational "
    "assistant (Mia) and an administrator panel for user and report management, "
    "and to test and evaluate the system. The study will benefit members of the "
    "public by providing a secure, accessible screening tool with personal report "
    "history, healthcare providers by offering a preliminary screening aid, and "
    "administrators by enabling oversight of registered users. The study adopted "
    "and was guided by the Technology Acceptance Model; data were gathered through "
    "interviews, questionnaires, document review, and observation, and the artifact "
    "was implemented as a Vue 3 single-page application communicating over a REST API "
    "with a Python/Django backend using PostgreSQL and an integrated large-language-"
    "model provider. The findings showed that the system accurately scores all three "
    "instruments, generates supportive bilingual (English and Swahili) feedback "
    "within a few seconds, links each assessment to the authenticated user\u2019s "
    "account, supports downloadable mental-health reports, and was rated easy to "
    "use during user-acceptance testing. The study recommends integrating a verified "
    "crisis-support pathway, expanding language and cultural adaptation, conducting "
    "larger clinical validation of AI-generated feedback, and strengthening role-based "
    "access control for administrators. The system demonstrates that a lightweight, "
    "authenticated, AI-assisted, and locally tailored platform can improve access "
    "to mental-health screening in low-resource settings."
)
page_break()

# ============================================================================= TOC
h1("TABLE OF CONTENTS")
add_toc()
page_break()

# ============================================================================= LIST OF TABLES
h1("LIST OF TABLES")
tables_list = [
    "Table 3.1: Structure of the self_assessment Table",
    "Table 3.2: Structure of the child_assessment Table",
    "Table 3.3: Structure of the model_config Table",
    "Table 3.4: Structure of the feature_model_assignment Table",
    "Table 3.7: Test Cases for the System",
]
for t in tables_list:
    para(t)
page_break()

# ============================================================================= LIST OF FIGURES
h1("LIST OF FIGURES")
figs = [
    "Figure 2.1: Screenshot of the BetterHelp Online Therapy Platform Interface",
    "Figure 2.2: Screenshot of the Wysa AI Mental-Health Chatbot Interface",
    "Figure 2.3: Screenshot of the MindDoc Mental-Health Self-Assessment Application",
    "Figure 2.4: Screenshot of the 7 Cups Online Emotional-Support Platform",
    "Figure 2.5: Screenshot of the Headspace Mental-Wellness Platform",
    "Figure 3.1: Use Case Diagram for the System",
    "Figure 3.2: Activity Diagram for the \u201cTake Assessment and Receive AI Feedback\u201d Process",
    "Figure 3.3: Sequence Diagram for the \u201cTake Assessment\u201d Interaction",
    "Figure 3.4: Class Diagram for the System",
    "Figure 3.5: Deployment Diagram Showing the System Architecture",
    "Figure 3.6: Entity Relationship Diagram (ERD) for the Database",
    "Figure 4.1: Home / Landing Screen of the System",
    "Figure 4.2: Mental Health Dashboard of the System",
    "Figure 4.3: Assessment Selection Screen",
    "Figure 4.4: PHQ-9 Assessment Question Screen",
    "Figure 4.5: Computed Score and AI-Generated Feedback Screen",
    "Figure 4.6: Conversational AI Assistant (Mia) Chat Screen",
    "Figure 4.7: Follow-up Referral Form",
    "Figure 4.8: Kiswahili (Swahili) Language Interface",
    "Figure 4.9: User Registration and Account Creation Screen",
    "Figure 4.10: Administrator Panel for User Management",
    "Figure 4.11: My Reports Screen Showing Assessment History",
    "Figure 4.12: Sample Downloaded Mental-Health Report (.docx)",
]
for f in figs:
    para(f)
page_break()

# ============================================================================= ACRONYMS
h1("LIST OF ACRONYMS AND ABBREVIATIONS")
abbr = [
    ("AI", "Artificial Intelligence"),
    ("API", "Application Programming Interface"),
    ("CBT", "Cognitive Behavioural Therapy"),
    ("CSS", "Cascading Style Sheets"),
    ("CORS", "Cross-Origin Resource Sharing"),
    ("DBMS", "Database Management System"),
    ("DRF", "Django REST Framework"),
    ("DSR", "Design Science Research"),
    ("ERD", "Entity Relationship Diagram"),
    ("GAD-7", "Generalized Anxiety Disorder Scale (7-item)"),
    ("HTML", "HyperText Markup Language"),
    ("HTTP", "HyperText Transfer Protocol"),
    ("ICT", "Information and Communication Technology"),
    ("JWT", "JSON Web Token"),
    ("LLM", "Large Language Model"),
    ("NLP", "Natural Language Processing"),
    ("ORM", "Object-Relational Mapping"),
    ("PHQ-9", "Patient Health Questionnaire (9-item)"),
    ("REST", "Representational State Transfer"),
    ("SDQ", "Strengths and Difficulties Questionnaire"),
    ("SPA", "Single-Page Application"),
    ("TAM", "Technology Acceptance Model"),
    ("UI", "User Interface"),
    ("UDSM", "University of Dar es Salaam"),
    ("URT", "United Republic of Tanzania"),
    ("WHO", "World Health Organization"),
    ("3NF", "Third Normal Form"),
]
for a, full in abbr:
    p = doc.add_paragraph()
    p.add_run(a + " \u2013 ").bold = True
    p.add_run(full)
page_break()

# ============================================================================= DEFINITION OF KEY TERMS
h1("DEFINITION OF KEY TERMS")
terms = [
    ("Self-Assessment", "the process by which a user evaluates their own mental-health "
     "status by answering a standardized questionnaire, without the direct "
     "involvement of a clinician."),
    ("User Account", "a registered account created through the system\u2019s "
     "registration page, secured with a username and password and authenticated "
     "via JSON Web Tokens (JWT)."),
    ("Administrator", "a privileged user (is_staff) who can manage registered "
     "accounts and download assessment reports for any user."),
    ("Mental Health Report", "a downloadable Word document (.docx) summarizing "
     "the user\u2019s assessment score, severity level, and AI-generated feedback."),
    ("PHQ-9", "the nine-item questionnaire used in this study to screen adults for "
     "symptoms of depression."),
    ("GAD-7", "the seven-item questionnaire used in this study to screen adults for "
     "symptoms of anxiety."),
    ("SDQ", "the behavioural screening questionnaire completed by a parent, guardian, "
     "or teacher to assess a child aged 4\u201317 years."),
    ("Large Language Model (LLM)", "the artificial-intelligence model used in this "
     "study to generate supportive, non-diagnostic feedback in the user\u2019s language."),
    ("Conversational Assistant (Mia)", "the artificial-intelligence chat feature of "
     "the system that provides empathetic emotional-support responses."),
    ("Follow-up Referral", "an optional form offered to users with elevated scores to "
     "voluntarily provide contact details so that support can be arranged."),
]
for term, meaning in terms:
    p = doc.add_paragraph()
    p.add_run(term + ": ").bold = True
    p.add_run(meaning)

# ----------------------------------------------- section break -> main document (arabic 1)
sec_main = doc.add_section(WD_SECTION.NEW_PAGE)
apply_margins(sec_main)
set_page_number_format(sec_main, "decimal", 1)
add_footer_page_number(sec_main)

# ============================================================================= CHAPTER ONE
h1("CHAPTER ONE: INTRODUCTION")

h2("1.1 Introduction")
para(
    "This chapter presents the background of the study, the statement of the problem, "
    "the objectives, the research questions, the significance of the study, its "
    "limitations and scope, and the organization of the study. It provides the "
    "foundation for understanding the development of an online mental-health "
    "self-assessment system and its role in improving access to mental-health "
    "screening, early detection, and guidance."
)

h2("1.2 Background of the Study")
para(
    "Mental health is an essential component of overall well-being and contributes "
    "significantly to productivity, social relationships, and quality of life. "
    "According to the World Health Organization (WHO, 2022), mental-health disorders "
    "such as depression, anxiety, and stress-related conditions are among the leading "
    "causes of disability worldwide, contributing to substantial social and economic "
    "burdens."
)
para(
    "Despite the rising prevalence of these challenges, access to services remains "
    "limited, particularly in developing countries, where there is a shortage of "
    "mental-health professionals, limited facilities, and widespread stigma (WHO, "
    "2021). Many individuals hesitate to seek help due to fear of discrimination and "
    "a lack of awareness about available services."
)
para(
    "Advances in information and communication technology (ICT) have transformed "
    "healthcare through digital-health solutions such as telemedicine, mobile-health "
    "applications, and web-based assessment tools (Torous et al., 2020). More "
    "recently, large language models (LLMs) have created new opportunities to deliver "
    "scalable, conversational, and personalized mental-health support, although their "
    "use also raises important safety and ethical considerations (Lawrence et al., "
    "2024)."
)
para(
    "In the Tanzanian context, mental-health services are still developing, and many "
    "individuals face barriers such as limited healthcare infrastructure, a shortage "
    "of professionals, and low awareness of mental-health conditions (URT, 2022). "
    "Although internet usage and smartphone adoption have increased, reliable, "
    "locally relevant, Swahili-capable web-based mental-health assessment platforms "
    "remain scarce. This situation creates a need for a digital platform that allows "
    "users to assess their mental-health status, receive feedback, and obtain "
    "appropriate recommendations for further support."
)

h2("1.3 Problem Statement")
para(
    "Mental-health disorders continue to affect many individuals worldwide, yet a "
    "significant proportion do not receive appropriate care due to limited resources, "
    "stigma, and inaccessible services (WHO, 2022). Traditional assessment methods "
    "often require physical visits to healthcare facilities, which may be costly, "
    "time-consuming, and inaccessible to people in remote areas, while social stigma "
    "further discourages help-seeking and leads to delayed detection (WHO, 2021). "
    "Although digital and AI-based tools have demonstrated value (Fitzpatrick et al., "
    "2017; Inkster et al., 2018), existing platforms are largely designed for "
    "high-income, English-speaking contexts and rarely provide a locally relevant, "
    "Swahili-capable, freely accessible solution that combines validated adult and "
    "child screening with immediate feedback and local referral (URT, 2022). There is "
    "therefore a need to design and implement an online mental-health self-assessment "
    "system that enables registered users to perform standardized self-assessments, "
    "receive immediate AI-assisted feedback in their own language, download personal "
    "mental-health reports, and optionally request follow-up support through the "
    "referral form, with administrators able to manage users and generate reports."
)

h2("1.4 Objectives of the Study")
h3("1.4.1 General Objective")
para(
    "To design and implement an online mental-health self-assessment system that "
    "supports early detection of mental-health concerns and delivers immediate, "
    "AI-assisted feedback and referral guidance for users in Tanzania."
)
h3("1.4.2 Specific Objectives")
roman_item("i", "To analyse user requirements and existing mental-health self-assessment practices.")
roman_item("ii", "To design the system architecture, database structure, and bilingual user interfaces.")
roman_item("iii", "To develop a web-based system that allows users to complete validated assessments (PHQ-9, GAD-7, and SDQ) and receive AI-generated feedback.")
roman_item("iv", "To integrate a conversational AI assistant (Mia) within the platform.")
roman_item("v", "To test and evaluate the functionality, accuracy, and usability of the developed system.")

h2("1.5 Research Questions")
para("Based on the specific objectives, the study was guided by the following research questions:")
roman_item("i", "What are the user requirements and existing practices for mental-health self-assessment in Tanzania?")
roman_item("ii", "How should the architecture, database, and bilingual interfaces of such a system be designed?")
roman_item("iii", "How can a web-based system enable validated assessments and AI-generated feedback?")
roman_item("iv", "How can a conversational AI assistant be integrated effectively within the platform?")
roman_item("v", "How well does the developed system perform in terms of functionality, accuracy, and usability?")

h2("1.6 Significance of the Study")
para(
    "This section explains who will benefit from the study and how. Questionnaire "
    "responses from 100 UDSM students informed the design of a system that provides "
    "registered users with a convenient, secure, and accessible platform for assessing "
    "their mental-health status, saving assessment history, and downloading personal "
    "reports; immediate AI-generated feedback will promote self-awareness and encourage "
    "timely help-seeking. Healthcare providers will "
    "benefit from a preliminary screening aid, and administrators will be able to "
    "manage user accounts and retrieve assessment reports for follow-up support."
)

h2("1.7 Limitations of the Study")
para(
    "The system relies on users\u2019 self-reported responses, which may be influenced "
    "by personal bias; this was mitigated by using internationally validated "
    "instruments (PHQ-9, GAD-7, and SDQ). Unstable internet connectivity, especially "
    "in rural areas, may restrict access, and the AI-feedback feature depends on an "
    "external service; this was mitigated by building a lightweight application "
    "optimized for low bandwidth and by returning the numerical score even when AI "
    "generation is delayed. Because feedback is produced by a large language model, "
    "there is a risk of generic or imperfect responses; the model is therefore "
    "restricted to a supportive, non-diagnostic role, and users with elevated scores "
    "are directed toward professional support (Lawrence et al., 2024)."
)

h2("1.8 Scope of the Study")
h3("1.8.1 Content Scope")
para(
    "The study covers requirements analysis, system design, development, and testing. "
    "The platform supports adults through the PHQ-9 and GAD-7 instruments and supports "
    "the screening of children aged 4\u201317 through the SDQ, completed by a parent, "
    "guardian, or teacher. It also includes AI-generated feedback, a conversational "
    "assistant. "
    "The system does not provide real-time counselling, professional therapy, or "
    "clinical diagnosis."
)
h3("1.8.2 Geographical Scope")
para(
    "The study focuses on the Tanzanian context, targeting members of the public, "
    "university students, parents, teachers, and mental-health professionals in "
    "Tanzania."
)
h3("1.8.3 Time Scope")
para(
    "The study was conducted during the 2025/2026 academic year, with requirements "
    "analysis, design, development, testing, and evaluation carried out within the "
    "allocated project timeframe of approximately six months."
)

h2("1.9 Organization of the Study")
para("Chapter One presents the introduction, including the background, problem statement, objectives, research questions, significance, limitations, and scope.")
para("Chapter Two reviews the theoretical literature, similar systems, the critical review and research gap, and a chapter summary.")
para("Chapter Three presents the methodology, including the research design, system development methodology, requirements analysis and modeling, database design, and testing design.")
para("Chapter Four presents the research findings and discussion, addressing each objective with the developed features and results.")
para("Chapter Five presents the summary of findings, conclusions, recommendations, and suggestions for further study.")

# ============================================================================= CHAPTER TWO
h1("CHAPTER TWO: LITERATURE REVIEW")

h2("2.1 Introduction")
para(
    "This chapter reviews the literature related to online and AI-assisted "
    "mental-health assessment systems. It presents the theoretical literature "
    "underpinning the study, reviews at least five similar systems from global to "
    "local contexts, provides a critical review that identifies the research gap, and "
    "concludes with a chapter summary."
)

h2("2.2 Theoretical Literature Review")
h3("2.2.1 The Technology Acceptance Model (TAM)")
para(
    "The study is guided by the Technology Acceptance Model (Davis, 1989), which "
    "explains technology adoption through two key constructs: perceived usefulness "
    "(the degree to which a user believes a system will enhance their performance or "
    "well-being) and perceived ease of use (the degree to which a user believes a "
    "system is free of effort). TAM helps explain the factors influencing users\u2019 "
    "acceptance of an anonymous, AI-assisted self-assessment tool."
)
h3("2.2.2 Design Science Research (DSR) Framework")
para(
    "The Design Science Research methodology (Hevner et al., 2004; Peffers et al., "
    "2007) guides the development of the system as an IT artifact through iterative "
    "design and evaluation cycles, which suit the complexity of building a "
    "mental-health system that must balance clinical validity, user experience, "
    "privacy, and technical performance."
)
h3("2.2.3 Concept of Mental-Health Assessment")
para(
    "Mental-health assessment is the systematic process of evaluating an "
    "individual\u2019s psychological, emotional, and behavioural functioning using "
    "standardized tools (American Psychiatric Association, 2013). It supports early "
    "detection and intervention, which help prevent severe complications and improve "
    "outcomes (Kazdin, 2000). Online assessment systems use validated tools such as "
    "the PHQ-9 for depression and the GAD-7 for anxiety (Kroenke et al., 2001; Spitzer "
    "et al., 2006), and, for children, the SDQ (Goodman, 1997)."
)
h3("2.2.4 Digital Mental-Health Technologies and Large Language Models")
para(
    "Digital mental-health technologies include mobile-health applications, web-based "
    "platforms, and artificial-intelligence systems that enable remote, anonymous, "
    "and confidential access, helping to reduce stigma (Torous et al., 2020; Naslund "
    "et al., 2017). Large language models can generate fluent, empathetic, and "
    "context-aware text for education, assessment support, and conversational "
    "intervention; however, the literature stresses risks such as factual "
    "inaccuracies, inconsistent responses, limited handling of crisis situations, and "
    "data-privacy concerns, and concludes that LLMs should serve as supportive aids "
    "rather than replacements for professional care (Lawrence et al., 2024). The "
    "present system adopts this position by using an LLM strictly to produce "
    "supportive, non-diagnostic feedback. Empirically, Fitzpatrick et al. (2017) "
    "showed in a randomized controlled trial that the Woebot chatbot significantly "
    "reduced depression symptoms, and Inkster et al. (2018) reported improvements in "
    "self-reported depression among engaged users of the Wysa chatbot."
)

h2("2.3 Similar Systems")
para(
    "The proposed system is not unique. This section reviews five existing digital "
    "mental-health platforms, from global to local relevance, and identifies how each "
    "differs from the proposed system. Because the proposed system is anonymous, "
    "bilingual (English and Swahili), free, and uses an LLM to generate supportive "
    "feedback, particular attention is given to AI-based and web-accessible platforms. "
    "Screenshots are included as evidence of the investigation of similar systems."
)

para("2.3.1 BetterHelp", bold=True)
para(
    "BetterHelp is one of the world\u2019s largest online therapy platforms, connecting "
    "users with licensed professionals through messaging, live chat, phone, and video "
    "sessions, with standardized screening during onboarding (BetterHelp, 2023). Its "
    "strengths include a large network of therapists, multi-modal communication, and "
    "web and mobile access. Its weaknesses include a high subscription cost that is "
    "inaccessible in low-income settings, no Swahili or local-language support, and a "
    "reliance on high-speed connectivity. Figure 2.1 shows the BetterHelp interface."
)
placeholder("Insert Figure 2.1: Screenshot of the BetterHelp platform (https://www.betterhelp.com)")
figure_caption("Figure 2.1: Screenshot of the BetterHelp Online Therapy Platform Interface")

para("2.3.2 Wysa", bold=True)
para(
    "Wysa is an empathy-driven, AI-based conversational agent for digital mental "
    "well-being that guides users through evidence-based techniques and incorporates "
    "the PHQ-9 (Inkster et al., 2018). Its strengths include empathetic AI "
    "interaction, a low-bandwidth text interface, and anonymous usage. Its weaknesses "
    "include no Swahili support and no local follow-up referral workflow tailored to Tanzania, "
    "and a focus on adults without a guardian/teacher child-screening pathway. Figure "
    "2.2 shows the Wysa interface."
)
placeholder("Insert Figure 2.2: Screenshot of the Wysa platform (https://www.wysa.com)")
figure_caption("Figure 2.2: Screenshot of the Wysa AI Mental-Health Chatbot Interface")

para("2.3.3 MindDoc", bold=True)
para(
    "MindDoc is a clinically developed application for mental-health self-monitoring "
    "through mood tracking and standardized questionnaires. Its strengths include a "
    "clinically grounded design, continuous symptom tracking, and referral "
    "recommendations. Its weaknesses include being predominantly app-based, locking "
    "core features behind a subscription, and a lack of localization for Sub-Saharan "
    "African languages and contexts. Figure 2.3 shows the MindDoc application."
)
placeholder("Insert Figure 2.3: Screenshot of the MindDoc application (https://www.minddoc.com)")
figure_caption("Figure 2.3: Screenshot of the MindDoc Mental-Health Self-Assessment Application")

para("2.3.4 7 Cups", bold=True)
para(
    "7 Cups is a web-based emotional-support platform connecting users with trained "
    "volunteer listeners and therapists, with anonymous chat, self-help guides, and "
    "assessments. Its strengths include anonymous peer support, free basic access, and "
    "community forums. Its weaknesses include variable volunteer quality, limited "
    "Swahili localization, and no AI-generated personalized feedback per assessment. "
    "Figure 2.4 shows the 7 Cups platform."
)
placeholder("Insert Figure 2.4: Screenshot of the 7 Cups platform (https://www.7cups.com)")
figure_caption("Figure 2.4: Screenshot of the 7 Cups Online Emotional-Support Platform")

para("2.3.5 Headspace", bold=True)
para(
    "Headspace is a digital wellness platform specializing in mindfulness-based "
    "interventions, guided meditation, and stress management. Its strengths include "
    "evidence-based programmes, high usability, and progress tracking. Its weaknesses "
    "include the absence of standardized clinical screening tools such as the PHQ-9 or "
    "GAD-7, a subscription cost, and a focus on prevention rather than screening and "
    "referral. Figure 2.5 shows the Headspace platform."
)
placeholder("Insert Figure 2.5: Screenshot of the Headspace platform (https://www.headspace.com)")
figure_caption("Figure 2.5: Screenshot of the Headspace Mental-Wellness Platform")

h2("2.4 Critical Review and Research Gap")
para(
    "A critical examination of the reviewed systems reveals limitations that this "
    "study addresses. Most platforms were developed for high-income, English-speaking "
    "contexts and lack adaptation to the linguistic, cultural, and infrastructural "
    "realities of Tanzania. Many are app-based or subscription-based, reducing "
    "accessibility; few combine validated adult and child screening within a single "
    "anonymous tool; and few connect users to local professional referrals. While AI "
    "chatbots such as Wysa demonstrate the value of conversational support, they are "
    "not Swahili-capable and do not integrate locally relevant referral information."
)
para(
    "The research gap, therefore, lies in the absence of a locally tailored, "
    "bilingual, freely accessible, AI-assisted online mental-health self-assessment "
    "platform that integrates validated adult and child screening, immediate "
    "non-diagnostic feedback for users in Tanzania. The innovation of the proposed "
    "innovation of the proposed system is its combination of (i) secure registered "
    "access with personal report history; (ii) three validated instruments (PHQ-9, "
    "GAD-7, and SDQ) in one platform; (iii) AI-generated, non-diagnostic feedback, "
    "downloadable reports, and a conversational assistant; and (iv) a bilingual "
    "English/Swahili interface with automatic language detection \u2014 a combination "
    "not offered by the reviewed systems."
)

h2("2.5 Chapter Summary")
para(
    "This chapter reviewed the theoretical literature, including the Technology "
    "Acceptance Model and the Design Science Research framework, the concept of "
    "mental-health assessment, and the emerging role of large language models. It "
    "reviewed five similar systems and, through a critical review, identified the "
    "research gap that the proposed system fills. The review establishes that "
    "locally tailored, bilingual, integrated, and AI-assisted solutions are needed in "
    "Tanzania and provides the foundation for the methodology presented in the next "
    "chapter."
)

# ============================================================================= CHAPTER THREE
h1("CHAPTER THREE: METHODOLOGY")

h2("3.1 Introduction")
para(
    "This chapter presents the research methodology and system design. It describes "
    "the research design, target population, sample size, and data-collection "
    "procedure and instruments; the system development methodology and its "
    "justification; the system requirements analysis and modeling; the database "
    "design; and the testing design."
)

h2("3.2 Research Design")
para(
    "The study adopted a descriptive and developmental research design and followed "
    "the Design Science Research (DSR) methodology, which focuses on creating and "
    "evaluating an innovative artifact \u2014 here, the online mental-health "
    "self-assessment system \u2014 to solve an identified problem. A mixed approach "
    "combining qualitative and quantitative methods was used: qualitative methods "
    "supported requirements gathering, while quantitative methods supported usability "
    "and performance evaluation."
)
h3("3.2.1 Target Population")
para(
    "The target population comprised undergraduate students at the University of Dar es "
    "Salaam (UDSM). This group was selected because university students frequently "
    "experience academic pressure, social transition, and financial stress, which "
    "increase their vulnerability to depression, anxiety, and related mental-health "
    "concerns. UDSM students also represent a digitally literate population that can "
    "realistically adopt a web-based self-assessment platform."
)
h3("3.2.2 Sample Size")
para(
    "The sample size was determined using Slovin\u2019s formula, which is applied when "
    "the population is large and a manageable sample is required while maintaining an "
    "acceptable margin of error. The formula is expressed as follows:"
)
para("n = N / (1 + Ne\u00b2)", bold=True)
para(
    "Where n is the required sample size, N is the total population size, and e is the "
    "margin of error expressed as a decimal. For this study, the estimated population "
    "of UDSM students was taken as N = 20,000, and a margin of error of e = 0.10 "
    "(10%) was adopted at a 95% confidence level. Substituting these values gives:"
)
para("n = 20,000 / (1 + 20,000 \u00d7 0.10\u00b2)", bold=True)
para("n = 20,000 / (1 + 200) = 20,000 / 201 \u2248 99.5", bold=True)
para(
    "The result was rounded to a sample of 100 UDSM students. Simple random sampling "
    "was used to select respondents from different colleges and programmes so that the "
    "sample reflected the diversity of the student population."
)
h3("3.2.3 Data Collection Procedure and Instruments")
para(
    "Data were collected primarily through questionnaires. This method was chosen "
    "because it allows standardized responses from a large number of participants "
    "within a limited time, supports quantitative analysis, and is appropriate for "
    "studies that seek to measure attitudes, perceptions, and usability preferences "
    "among student users."
)
para("Questionnaires: ", bold=True)
para(
    "A structured questionnaire was administered to the 100 UDSM students. The "
    "instrument contained closed-ended items using Likert-scale responses (for example, "
    "Strongly Disagree to Strongly Agree) and a small number of open-ended questions. "
    "The questionnaire gathered information on students\u2019 awareness of mental-health "
    "issues, willingness to use an online screening tool, preferred language "
    "(English or Kiswahili), expectations of AI-generated feedback, and perceived "
    "ease of use after interacting with the developed prototype. The completed "
    "questionnaires were checked for completeness before analysis."
)
para("Document Review: ", bold=True)
para(
    "A systematic review of academic publications, WHO reports, and the official "
    "documentation of the PHQ-9, GAD-7, and SDQ instruments informed the design of "
    "the screening modules and the interpretation of scores."
)

h2("3.3 System Development Methodology")
para(
    "The system was developed using the Design Science Research methodology, which "
    "comprises problem identification and motivation, definition of objectives, design "
    "and development, demonstration, evaluation, and communication of the results."
)
h3("3.3.1 Methodology Justification")
para(
    "DSR was chosen over alternatives such as Waterfall, Agile, and prototyping-only "
    "approaches because it is problem-driven and therefore well-suited to solving the "
    "clearly defined problem of limited access to mental-health screening in Tanzania; "
    "it promotes iterative design and evaluation, allowing continuous refinement based "
    "on stakeholder feedback; it is a rigorous, recognized approach in "
    "information-systems research that provides academic credibility (Hevner et al., "
    "2004); and it integrates both design and evaluation within a single framework, "
    "enabling the researcher to build the system and assess its effectiveness within "
    "the same methodological structure."
)
h3("3.3.2 Development Tools and Technologies")
bullet("Frontend: HTML5, CSS3, JavaScript, and the Vue 3 framework, built with Vite, using Vue Router, Pinia (authentication state), vue-i18n (English/Swahili), and Axios with JWT interceptors.")
bullet("Backend: Python with the Django framework and the Django REST Framework (DRF), with djangorestframework-simplejwt for authentication, drf-yasg for API documentation, and langdetect for language detection.")
bullet("Artificial Intelligence: a hosted large language model accessed through an OpenAI-compatible client (DeepSeek / OpenAI), configurable per feature.")
bullet("Database: SQLite for development and PostgreSQL for production, accessed via the Django ORM.")
bullet("Server: Gunicorn application server behind an Nginx web server.")
bullet("Design tools: draw.io and Lucidchart for the UML and ER diagrams.")

h2("3.4 System Requirements Analysis")
para(
    "Requirements analysis identifies the functional and non-functional requirements "
    "the system must satisfy. The system has two primary roles: a Registered User "
    "(a member of the public who creates an account and completes assessments) and "
    "an Administrator (a staff user who manages accounts and generates user reports). "
    "Parents, guardians, and teachers complete the SDQ on behalf of a child."
)
h3("3.4.1 Functional Requirements")
para("User Registration and Login: ", bold=True)
para("The system shall allow new users to register with a username, email, and password, and shall authenticate users via JSON Web Tokens (JWT) before granting access to protected features.")
para("Adult Self-Assessment (PHQ-9 and GAD-7): ", bold=True)
para("The system shall present the selected instrument sequentially, collect the user\u2019s age group and sex, record responses, compute the total score, and store the result linked to the authenticated user.")
para("Child Assessment (SDQ): ", bold=True)
para("The system shall allow a parent, guardian, or teacher to complete the SDQ for a child aged 4\u201317 and shall compute the total difficulties and prosocial scores.")
para("AI-Generated Feedback: ", bold=True)
para("On completion, the system shall generate immediate, personalized, supportive, and non-diagnostic feedback using a large language model, and shall persist the feedback text with the assessment record.")
para("Downloadable Mental-Health Report: ", bold=True)
para("The system shall allow a registered user to download a Word document (.docx) summarizing the assessment score, severity level, and AI-generated feedback; administrators shall be able to download reports for any user.")
para("Conversational AI Assistant (Mia): ", bold=True)
para("The system shall provide an AI chat assistant for emotional support that maintains short conversational context.")
para("Language Detection and Bilingual Interface: ", bold=True)
para("The system shall provide an English/Swahili interface and detect the language of the user\u2019s input to respond appropriately.")
para("Follow-up Referral: ", bold=True)
para("When a score indicates elevated concern, the system shall offer a follow-up form to voluntarily collect contact details, location, and (for children) school.")
para("Administrator User Management: ", bold=True)
para("The system shall allow an administrator to view all registered users, add new users, remove users, and access each user\u2019s assessment history.")
h3("3.4.2 Non-Functional Requirements")
para("Security and Authentication: ", bold=True)
para("User passwords shall be hashed; API access shall require valid JWT tokens; administrator endpoints shall be restricted to staff users; and data shall be transmitted over HTTPS with restricted CORS origins.")
para("Privacy: ", bold=True)
para("Assessment demographic data (age group and sex) shall be collected without requiring real names; additional contact details shall be collected only when a user voluntarily requests follow-up support.")
para("Security: ", bold=True)
para("Data shall be transmitted over HTTPS, with restricted CORS origins and CSRF protections, and all input shall be validated.")
para("Usability: ", bold=True)
para("The bilingual interface shall be simple, intuitive, and empathetic, enabling users of all skill levels to complete tasks with minimal effort.")
para("Performance: ", bold=True)
para("Scores shall be available immediately, and AI-generated feedback shall typically be returned within a few seconds, depending on the provider.")
para("Reliability and Scalability: ", bold=True)
para("Scoring shall be accurate, the system shall degrade gracefully if the external AI service is unavailable, and the modular architecture shall support a growing number of users.")
h3("3.4.3 Modeling Language")
para(
    "The Unified Modeling Language (UML) was used to design the system before "
    "development, clarifying how the system works and how users interact with it. The "
    "following models were produced."
)
h3("3.4.3.1 Use Case Diagram")
para(
    "Figure 3.1 presents the use case diagram. The primary actors are the Registered "
    "User and the Administrator, with the Parent/Guardian acting on behalf of a child "
    "and an external AI Provider as a supporting actor. Use cases include "
    "registration and login, taking the PHQ-9, GAD-7, and SDQ assessments, receiving "
    "AI-generated feedback, downloading a mental-health report, chatting with the AI "
    "assistant, submitting a follow-up referral, and (for administrators) managing "
    "users and generating user reports."
)
add_diagram("usecase.png")
figure_caption("Figure 3.1: Use Case Diagram for the System")
h3("3.4.3.2 Activity Diagram")
para(
    "Figure 3.2 presents the activity diagram for taking an assessment and receiving "
    "AI feedback. The user selects an assessment, provides age group and sex, answers "
    "each item, and on completion the system computes the score, stores a user-linked "
    "record with the AI feedback, displays the result, offers report download, and "
    "if the score indicates elevated concern, offers the follow-up referral form."
)
add_diagram("activity.png", width_inches=4.0)
figure_caption("Figure 3.2: Activity Diagram for the \u201cTake Assessment and Receive AI Feedback\u201d Process")
h3("3.4.3.3 Sequence Diagram")
para(
    "Figure 3.3 presents the sequence diagram, which shows the message exchange among "
    "the User, the Vue user interface, the Django REST API (with JWT authentication), "
    "the database, and the external AI provider when an assessment is taken."
)
add_diagram("sequence.png")
figure_caption("Figure 3.3: Sequence Diagram for the \u201cTake Assessment\u201d Interaction")
h3("3.4.3.4 Class Diagram")
para(
    "Figure 3.4 presents the class diagram. The main classes are User (Django auth), "
    "SelfAssessment, ChildAssessment, ModelConfig, and "
    "FeatureModelAssignment. Each assessment record is linked to a User through a "
    "foreign key and stores the AI feedback text and severity level."
)
add_diagram("classes.png")
figure_caption("Figure 3.4: Class Diagram for the System")
h3("3.4.3.5 Deployment Diagram")
para(
    "Figure 3.5 presents the deployment diagram, showing the client device (a browser "
    "running the Vue single-page application), the application server (Nginx, "
    "Gunicorn, and Django), the PostgreSQL database, and the external AI provider, "
    "with HTTPS communication between the tiers."
)
add_diagram("deployment.png")
figure_caption("Figure 3.5: Deployment Diagram Showing the System Architecture")

h2("3.5 Database Design")
para(
    "The database was designed to ensure data integrity and efficient retrieval and "
    "is implemented using a relational DBMS \u2014 SQLite during development and "
    "PostgreSQL in production \u2014 accessed through the Django ORM. The database "
    "includes Django\u2019s built-in auth_user table for registered accounts and "
    "assessment tables linked to users through foreign keys. Assessment questions "
    "are defined in the application rather than stored as database records."
)
h3("3.5.1 Table Design")
para("Table 3.1 shows the structure of the self_assessment table, which stores adult PHQ-9 and GAD-7 records.")
caption("Table 3.1: Structure of the self_assessment Table")
simple_table(
    ["Field", "Data Type", "Key", "Description"],
    [
        ["id", "INT", "PK", "Unique identifier (auto-increment)"],
        ["user_id", "INT", "FK", "Reference to auth_user.id"],
        ["assessment_type", "VARCHAR(10)", "", "Instrument used (phq9 or gad7)"],
        ["age_group", "VARCHAR(10)", "", "User age band (e.g., 16\u201320)"],
        ["sex", "VARCHAR(10)", "", "Male / Female"],
        ["score", "INT", "", "Computed total score"],
        ["ai_response", "TEXT", "", "Persisted AI-generated feedback"],
        ["severity_level", "VARCHAR(30)", "", "Derived severity label"],
        ["contact_info", "VARCHAR(255)", "", "Optional, from follow-up form"],
        ["location", "VARCHAR(255)", "", "Optional, from follow-up form"],
        ["description", "TEXT", "", "Optional notes from follow-up form"],
        ["created_at", "TIMESTAMP", "", "Time the record was created"],
    ],
)
para("Table 3.2 shows the structure of the child_assessment table, which stores SDQ records.")
caption("Table 3.2: Structure of the child_assessment Table")
simple_table(
    ["Field", "Data Type", "Key", "Description"],
    [
        ["id", "INT", "PK", "Unique identifier (auto-increment)"],
        ["user_id", "INT", "FK", "Reference to auth_user.id"],
        ["assessment_type", "VARCHAR(10)", "", "Instrument used (sdq)"],
        ["age_group", "VARCHAR(10)", "", "Child age band (4\u20137, 8\u201312, 13\u201317)"],
        ["sex", "VARCHAR(10)", "", "Male / Female"],
        ["difficulties", "INT", "", "Total difficulties score"],
        ["prosocial", "INT", "", "Prosocial behaviour score"],
        ["score", "INT", "", "Computed total score"],
        ["ai_response", "TEXT", "", "Persisted AI-generated feedback"],
        ["severity_level", "VARCHAR(30)", "", "Derived severity label"],
        ["contact_info", "VARCHAR(255)", "", "Optional, from follow-up form"],
        ["school", "VARCHAR(255)", "", "Optional, from follow-up form"],
        ["location", "VARCHAR(255)", "", "Optional, from follow-up form"],
        ["description", "TEXT", "", "Optional notes from follow-up form"],
        ["created_at", "TIMESTAMP", "", "Time the record was created"],
    ],
)
para("Table 3.3 shows the structure of the model_config table, which stores AI model configuration.")
caption("Table 3.3: Structure of the model_config Table")
simple_table(
    ["Field", "Data Type", "Key", "Description"],
    [
        ["id", "INT", "PK", "Unique identifier"],
        ["name", "VARCHAR(50)", "", "Model name (e.g., deepseek-chat)"],
        ["provider", "VARCHAR(20)", "", "Provider (openai / deepseek)"],
        ["base_url", "VARCHAR", "", "Optional API base URL"],
        ["temperature", "FLOAT", "", "Sampling temperature"],
        ["active", "BOOLEAN", "", "Whether the model is active"],
    ],
)
para("Table 3.4 shows the structure of the feature_model_assignment table, which maps a feature to a model.")
caption("Table 3.4: Structure of the feature_model_assignment Table")
simple_table(
    ["Field", "Data Type", "Key", "Description"],
    [
        ["id", "INT", "PK", "Unique identifier"],
        ["feature_key", "VARCHAR(50)", "UNIQUE", "Feature (e.g., phq9, gad7, sdq)"],
        ["model_id", "INT", "FK", "References model_config.id"],
    ],
)
h3("3.5.2 Third Normal Form (3NF)")
para(
    "The schema was normalized to the Third Normal Form. Every table has a primary "
    "key, all non-key attributes depend only on the primary key, and there are no "
    "transitive dependencies. For example, AI model details are stored once in the "
    "model_config table and referenced from feature_model_assignment through the "
    "model_id foreign key rather than being duplicated. Figure 3.6 presents the entity "
    "relationship diagram for the normalized database."
)
add_diagram("erd.png")
figure_caption("Figure 3.6: Entity Relationship Diagram (ERD) for the Database")

h2("3.6 Testing Design")
para(
    "Multiple forms of testing were adopted to justify the usability and correctness "
    "of the system. Unit testing verified individual components such as score "
    "computation, language detection, and input validation; integration testing "
    "verified the end-to-end flow across the user interface, the API, the database, "
    "and the AI provider; user-acceptance testing involved students, parents/teachers, "
    "and professionals performing real tasks; security testing verified HTTPS, CORS, "
    "and CSRF protections; and performance testing measured response and "
    "feedback-generation times."
)
h3("3.6.1 Test Cases")
para("Table 3.7 shows the test cases designed for the system.")
caption("Table 3.7: Test Cases for the System")
simple_table(
    ["Test Case", "Scope", "Method", "Expected Outcome"],
    [
        ["TC-01 Scoring", "PHQ-9 / GAD-7 / SDQ scoring", "Submit known responses", "Correct total score computed"],
        ["TC-02 Validation", "Incomplete submissions", "Submit missing answers", "Request rejected with an error"],
        ["TC-03 Language", "English & Swahili input", "Submit text in each language", "Correct language detected; feedback in that language"],
        ["TC-04 AI Feedback", "Feedback generation", "Complete an assessment", "Supportive, non-diagnostic feedback returned"],
        ["TC-05 Referral", "Elevated score", "Submit a high-scoring assessment", "Follow-up referral form offered"],
        ["TC-06 Security", "Transport & access", "Cross-origin / invalid request", "Request blocked; HTTPS enforced"],
        ["TC-07 Performance", "Response time", "Normal usage load", "Score immediate; feedback within a few seconds"],
        ["TC-08 Registration", "User sign-up", "Submit valid registration form", "Account created; user can log in"],
        ["TC-09 My Reports", "Report history", "Complete assessment while logged in", "Report listed; .docx downloads successfully"],
        ["TC-10 Admin", "User/report management", "Admin selects user in panel", "User reports listed; admin can download .docx"],
    ],
)

h2("3.7 Chapter Summary")
para(
    "This chapter presented the research design, the UDSM student target population, "
    "the sample size of 100 respondents determined through Slovin\u2019s formula, and "
    "the use of questionnaires as the primary data-collection instrument; justified "
    "the Design Science Research methodology and listed the development tools; "
    "documented the functional and non-functional requirements and the UML models; "
    "presented the normalized database design and its four tables; and described the "
    "testing design and test cases. These artifacts provide a well-documented "
    "foundation for the implementation and findings presented in the next chapter."
)

# ============================================================================= CHAPTER FOUR
h1("CHAPTER FOUR: RESEARCH FINDINGS AND DISCUSSION")

h2("4.1 Introduction")
para(
    "This chapter presents the findings of the study and discusses them in relation to "
    "each specific objective. It describes the developed system, supported by "
    "screenshots, and the results of testing and evaluation."
)

h2("4.2 Presentation of Findings")

h3("4.2.1 Objective One: Requirements Analysis")
para(
    "The requirements analysis, conducted through questionnaires administered to 100 "
    "UDSM students together with document review of validated screening instruments, "
    "confirmed a need for a secure, bilingual, easy-to-use screening tool that "
    "provides immediate AI-generated feedback, stores personal assessment history, "
    "and supports downloadable reports. These findings informed the functional and "
    "non-functional requirements presented in Section 3.4."
)

h3("4.2.2 Objective Two: System and Database Design")
para(
    "The architecture, database, and bilingual interfaces were designed using the UML "
    "models in Section 3.4.3 and the normalized database in Section 3.5. The result is "
    "a decoupled client\u2013server architecture in which a Vue single-page application "
    "communicates over a REST API with a Django backend and a PostgreSQL database, "
    "with an external large-language-model provider for feedback generation."
)

h3("4.2.3 Objective Three: Development of Assessments, Authentication, and AI Feedback")
para(
    "On first visit, the user is presented with the public landing screen shown in "
    "Figure 4.1, from which the Mental Health entry point leads to registration and "
    "login. New users create an account through the registration screen in Figure 4.9, "
    "entering a full name, username, email, and password; returning users authenticate "
    "with their credentials before accessing protected features. After login, the user "
    "is redirected to the dashboard in Figure 4.2, which displays a personalised "
    "welcome message (for example, \u201cWelcome, munirath\u201d or \u201cKaribu, "
    "munirath\u201d in Kiswahili). Each completed assessment is stored against the "
    "authenticated user\u2019s account."
)
add_screenshot("4_1_landing.png")
figure_caption("Figure 4.1: Home / Landing Screen of the System")
add_screenshot("4_9_register.png")
figure_caption("Figure 4.9: User Registration and Account Creation Screen")
add_screenshot("4_2_dashboard.png")
figure_caption("Figure 4.2: Mental Health Dashboard of the System")
para(
    "From the dashboard the user selects \u201cAssess Yourself\u201d and chooses one of "
    "the three instruments, as shown in the assessment selection screen in Figure 4.3. "
    "After providing age group and sex, the user answers the standardized items one at "
    "a time with a progress indicator, as shown for the PHQ-9 in Figure 4.4. Once all "
    "items are answered, the responses are sent to the backend, the total score is "
    "computed against the instrument range, and a large language model generates "
    "supportive, non-diagnostic feedback in the user\u2019s language; the resulting "
    "score and feedback are shown in Figure 4.5."
)
add_screenshot("4_3_selection.png")
figure_caption("Figure 4.3: Assessment Selection Screen")
add_screenshot("4_4_question.png")
figure_caption("Figure 4.4: PHQ-9 Assessment Question Screen")
add_screenshot("4_5_feedback.png")
figure_caption("Figure 4.5: Computed Score and AI-Generated Feedback Screen")

h3("4.2.4 Objective Four: Integration of the Assistant, Referral, and Bilingual Support")
para(
    "The conversational AI assistant (Mia) was integrated and is available from the "
    "dashboard and from a floating action button on every screen. As shown in Figure "
    "4.6, Mia greets the user, maintains recent conversation context, and responds "
    "with empathetic, markdown-formatted guidance while remaining strictly "
    "non-diagnostic. When an assessment score reaches the elevated-risk threshold, the "
    "user is offered the follow-up referral form shown in Figure 4.7, which allows the "
    "user to leave optional contact details so that local professional support can be "
    "arranged. Bilingual support was implemented throughout the interface using "
    "vue-i18n; Figure 4.8 shows the same dashboard rendered in Kiswahili after the "
    "language is switched, confirming that all user-facing strings are translated."
)
add_screenshot("4_6_chat.png")
figure_caption("Figure 4.6: Conversational AI Assistant (Mia) Chat Screen")
add_screenshot("4_7_followup.png")
figure_caption("Figure 4.7: Follow-up Referral Form")
add_screenshot("4_8_swahili.png")
figure_caption("Figure 4.8: Kiswahili (Swahili) Language Interface")

h3("4.2.5 Personal Reports, Downloadable Documents, and Administrator Panel")
para(
    "After completing an assessment, the user can view all saved results on the "
    "\u201cMy Reports\u201d page shown in Figure 4.11. Each report card displays the "
    "instrument name, score, severity level, and date completed. The user can download "
    "a Word document (.docx) summarising the assessment; Figure 4.12 shows a sample "
    "downloaded report containing the user name, assessment type, total score, severity "
    "level, demographic details, and personalised AI feedback."
)
add_screenshot("4_11_my_reports.png")
figure_caption("Figure 4.11: My Reports Screen Showing Assessment History")
add_screenshot("4_12_report_docx.png")
figure_caption("Figure 4.12: Sample Downloaded Mental-Health Report (.docx)")
para(
    "Staff users with administrator privileges access the Admin Panel shown in "
    "Figure 4.10 to manage registered accounts. The panel allows an administrator to "
    "add new users (with optional staff status), remove existing users, select any "
    "registered user, view that user\u2019s assessment history, and download the same "
    ".docx reports on the user\u2019s behalf for follow-up or institutional support."
)
add_screenshot("4_10_admin.png")
figure_caption("Figure 4.10: Administrator Panel for User Management")

h3("4.2.6 Objective Five: Testing and Evaluation")
para(
    "The test cases in Table 3.7 were executed. Unit tests confirmed accurate scoring "
    "for all three instruments and correct handling of incomplete submissions; "
    "integration tests verified the end-to-end flow to the AI provider and back; "
    "security tests confirmed HTTPS, restricted cross-origin access, and input "
    "validation; and performance tests showed that scores were returned immediately "
    "while AI feedback was typically returned within a few seconds. During "
    "user-acceptance testing among UDSM students, participants found the system easy to "
    "use, valued the immediate bilingual feedback, appreciated the ability to save and "
    "download personal reports, and rated the registration and dashboard experience "
    "positively."
)
para(
    "These findings indicate that the system meets its objectives. Consistent with the "
    "literature (Lawrence et al., 2024), the AI component is deliberately constrained "
    "to a supportive, non-diagnostic role, and elevated-risk users are directed toward "
    "professional support."
)

h2("4.3 Chapter Summary")
para(
    "This chapter presented the findings for each objective, supported by screenshots "
    "of the developed system and the results of testing and evaluation. The findings "
    "confirm that the implemented platform is an authenticated, bilingual, AI-assisted "
    "self-assessment system with validated adult and child instruments, personal "
    "report history, downloadable Word reports, administrator user management, and a "
    "conversational assistant."
)

# ============================================================================= CHAPTER FIVE
h1("CHAPTER FIVE: SUMMARY, CONCLUSIONS AND RECOMMENDATIONS")

h2("5.1 Introduction")
para(
    "This chapter summarizes the findings of the study in relation to each objective, "
    "draws conclusions in response to the research questions, and provides "
    "recommendations and suggestions for further study."
)

h2("5.2 Summary of Findings")
para(
    "The study designed and implemented MindCare, an online mental-health "
    "self-assessment system for UDSM students and the wider Tanzanian context. Data "
    "were collected from 100 UDSM students using questionnaires, and the sample size "
    "was justified through Slovin\u2019s formula. The key findings are summarised below "
    "by objective."
)
h3("5.2.1 Objective One: Requirements Analysis")
para(
    "Analysis of questionnaire responses from 100 UDSM students revealed strong demand "
    "for a digital mental-health screening tool that is accessible on ordinary "
    "smartphones and laptops, available in both English and Kiswahili, and capable of "
    "returning immediate feedback without requiring a clinic visit. Students reported "
    "that stigma and time constraints often prevent them from seeking face-to-face "
    "counselling, but they were willing to complete validated online questionnaires "
    "when the process was simple and private. The findings also showed that users "
    "expected to create a personal account, review past assessments, and download a "
    "written summary of their results. These requirements directly shaped the "
    "functional specifications documented in Chapter Three."
)
h3("5.2.2 Objective Two: System and Database Design")
para(
    "A decoupled three-tier architecture was designed and documented using UML "
    "diagrams and an entity relationship diagram. The Vue 3 single-page application "
    "handles the user interface and bilingual presentation; the Django REST backend "
    "enforces JWT authentication, computes scores, stores user-linked assessment "
    "records, and communicates with an external large-language-model provider; and "
    "the PostgreSQL database (SQLite in development) persists users, assessments, AI "
    "feedback text, and severity labels. The normalized schema separates adult "
    "(self_assessment) and child (child_assessment) records, links every assessment "
    "to auth_user through a foreign key, and supports downloadable report generation "
    "without duplicating AI model configuration data."
)
h3("5.2.3 Objective Three: Development of Assessments and AI Feedback")
para(
    "The implemented system allows registered users to complete the PHQ-9, GAD-7, and "
    "SDQ instruments and receive immediate AI-generated feedback. As shown in Figures "
    "4.1, 4.9, and 4.2, users register, log in, and access a personalised dashboard "
    "before starting an assessment. Figure 4.3 to Figure 4.5 demonstrate the full "
    "assessment flow from instrument selection through item presentation to computed "
    "score and Mia\u2019s supportive response. Scores are calculated deterministically "
    "on the server, stored with the authenticated user, and mapped to severity labels "
    "such as Minimal, Moderate, or Moderately Severe for PHQ-9. The AI component "
    "remains non-diagnostic and generates empathetic guidance in the detected language."
)
h3("5.2.4 Objective Four: Integration of the Conversational Assistant and Bilingual Support")
para(
    "The Mia conversational assistant was integrated as a persistent support channel "
    "available from the dashboard and through a floating action button. Figure 4.6 "
    "shows the chat interface, which maintains short conversational context and "
    "returns markdown-formatted responses. Bilingual support was implemented with "
    "vue-i18n across all major screens; Figure 4.8 confirms that navigation labels, "
    "welcome messages, and dashboard text switch correctly to Kiswahili. When an "
    "assessment score indicates elevated concern, the follow-up referral form in "
    "Figure 4.7 allows the user to volunteer contact details for further support."
)
h3("5.2.5 Objective Five: Testing, Reports, and Evaluation")
para(
    "Functional testing confirmed accurate scoring, secure JWT-protected routes, and "
    "successful end-to-end communication with the AI provider. Figures 4.11 and 4.12 "
    "demonstrate the report-management features: users can review assessment history "
    "on the My Reports page and download a structured Word document containing the "
    "score, severity, demographics, and personalised feedback. Figure 4.10 shows the "
    "administrator panel, through which staff users add or remove accounts and access "
    "any user\u2019s reports for institutional follow-up. User-acceptance testing with "
    "UDSM students indicated that the registration flow, dashboard layout, report "
    "download, and bilingual interface were easy to understand and met the "
    "expectations expressed in the questionnaire study."
)

h2("5.3 Conclusions")
para(
    "In response to the general objective, the study successfully designed and "
    "implemented an online mental-health self-assessment system that supports early "
    "awareness of depression, anxiety, and child behavioural concerns among UDSM "
    "students and similar user groups in Tanzania."
)
para(
    "With respect to the specific objectives, the following conclusions are drawn. "
    "First, questionnaire data from 100 UDSM students confirmed that there is a "
    "documented need for a secure, bilingual, web-based screening tool with personal "
    "report history. Second, the system architecture and database design provide a "
    "scalable foundation for authenticated multi-user operation. Third, the PHQ-9, "
    "GAD-7, and SDQ modules operate correctly and produce immediate AI-assisted "
    "feedback linked to each user account. Fourth, the Mia assistant and Kiswahili "
    "interface improve accessibility and user comfort. Fifth, testing and evaluation "
    "showed that the platform is usable, that downloadable reports add practical "
    "value for students and counsellors, and that the administrator panel supports "
    "basic user and report management."
)
para(
    "Overall, the study concludes that a lightweight, authenticated, bilingual, "
    "AI-assisted web platform can improve access to mental-health screening in a "
    "low-resource university setting. By combining validated instruments with "
    "supportive feedback, personal report storage, and optional follow-up referral, "
    "MindCare addresses gaps of accessibility, integration, stigma reduction, and "
    "localization identified in the literature, while keeping the AI component within "
    "a responsible, non-diagnostic role."
)

h2("5.4 Recommendations")
para(
    "Based on the findings and conclusions, the following recommendations are offered "
    "for improving and sustaining the system."
)
roman_item("i", "A verified crisis-support pathway with local emergency numbers and university counselling contacts should be integrated so that users with severe scores receive immediate guidance on where to seek help.")
roman_item("ii", "UDSM health services and student-affairs offices should pilot the platform as a supplementary screening and awareness tool, with clear referral protocols for students whose results indicate elevated concern.")
roman_item("iii", "The Kiswahili feedback and interface text should be reviewed by bilingual mental-health professionals to strengthen cultural relevance and readability.")
roman_item("iv", "A follow-up study with a larger sample and clinical supervision should evaluate the quality, safety, and consistency of AI-generated feedback over time.")
roman_item("v", "Privacy-preserving, aggregated analytics should be added so that institutions can monitor general usage trends without exposing individual responses.")

h2("5.5 Suggestions for Further Study")
para(
    "Further research may explore fine-tuning or domain-adapting the language model on "
    "vetted Swahili mental-health content, adding voice-based interaction for users "
    "with low literacy, and developing secure, consent-based pathways for professional "
    "follow-up support. Offline or progressive-web-app capabilities could also be "
    "investigated to improve access under intermittent connectivity. Longitudinal "
    "studies could examine whether repeated use of the platform improves help-seeking "
    "behaviour among UDSM students over an academic year."
)

h2("5.6 Chapter Summary")
para(
    "This chapter summarised the findings by objective using evidence from the "
    "questionnaire study of 100 UDSM students and from the implemented MindCare "
    "system, including registration, assessment, My Reports, downloadable documents, "
    "and administrator functions. It presented conclusions in response to the "
    "research questions, offered recommendations for deployment and improvement, and "
    "suggested directions for further study, thereby completing the report."
)

# ============================================================================= REFERENCES
h1("REFERENCES")
para(
    "All references are presented in accordance with the American Psychological "
    "Association (APA) 7th Edition citation style."
)
refs = [
    "American Psychiatric Association. (2013). Diagnostic and statistical manual of mental disorders (5th ed.). American Psychiatric Publishing.",
    "Andersson, G., & Titov, N. (2014). Advantages and limitations of Internet-based interventions for common mental disorders. World Psychiatry, 13(1), 4\u201311. https://doi.org/10.1002/wps.20083",
    "Barak, A., Hen, L., Boniel-Nissim, M., & Shapira, N. (2008). A comprehensive review and a meta-analysis of the effectiveness of Internet-based psychotherapeutic interventions. Journal of Technology in Human Services, 26(2\u20134), 109\u2013160. https://doi.org/10.1080/15228830802094429",
    "Bennett, G. G., & Glasgow, R. E. (2009). The delivery of public health interventions via the Internet: Actualizing their potential. Annual Review of Public Health, 30, 273\u2013292. https://doi.org/10.1146/annurev.publhealth.031308.100235",
    "BetterHelp. (2023). BetterHelp: Professional therapy with a licensed therapist. https://www.betterhelp.com",
    "Davis, F. D. (1989). Perceived usefulness, perceived ease of use, and user acceptance of information technology. MIS Quarterly, 13(3), 319\u2013340. https://doi.org/10.2307/249008",
    "Fitzpatrick, K. K., Darcy, A., & Vierhile, M. (2017). Delivering cognitive behavior therapy to young adults with symptoms of depression and anxiety using a fully automated conversational agent (Woebot): A randomized controlled trial. JMIR Mental Health, 4(2), e19. https://doi.org/10.2196/mental.7785",
    "Goodman, R. (1997). The Strengths and Difficulties Questionnaire: A research note. Journal of Child Psychology and Psychiatry, 38(5), 581\u2013586. https://doi.org/10.1111/j.1469-7610.1997.tb01545.x",
    "Headspace. (2023). Headspace: Meditation and sleep made simple. https://www.headspace.com",
    "Hevner, A. R., March, S. T., Park, J., & Ram, S. (2004). Design science in information systems research. MIS Quarterly, 28(1), 75\u2013105. https://doi.org/10.2307/25148625",
    "Inkster, B., Sarda, S., & Subramanian, V. (2018). An empathy-driven, conversational artificial intelligence agent (Wysa) for digital mental well-being: Real-world data evaluation mixed-methods study. JMIR mHealth and uHealth, 6(11), e12106. https://doi.org/10.2196/12106",
    "Joinson, A. N. (2001). Self-disclosure in computer-mediated communication: The role of self-awareness and visual anonymity. European Journal of Social Psychology, 31(2), 177\u2013192. https://doi.org/10.1002/ejsp.36",
    "Kazdin, A. E. (2000). Psychotherapy for children and adolescents: Directions for research and practice. Oxford University Press.",
    "Kroenke, K., Spitzer, R. L., & Williams, J. B. W. (2001). The PHQ-9: Validity of a brief depression severity measure. Journal of General Internal Medicine, 16(9), 606\u2013613. https://doi.org/10.1046/j.1525-1497.2001.016009606.x",
    "Lawrence, H. R., Schneider, R. A., Rubin, S. B., Matari\u0107, M. J., McDuff, D. J., & Jones Bell, M. (2024). The opportunities and risks of large language models in mental health. JMIR Mental Health, 11, e59479. https://doi.org/10.2196/59479",
    "Naslund, J. A., Aschbrenner, K. A., Araya, R., Marsch, L. A., Un\u00fctzer, J., Patel, V., & Bartels, S. J. (2017). Digital technology for treating and preventing mental disorders in low-income and middle-income countries: A narrative review of the literature. The Lancet Psychiatry, 4(6), 486\u2013500. https://doi.org/10.1016/S2215-0366(17)30096-2",
    "Peffers, K., Tuunanen, T., Rothenberger, M. A., & Chatterjee, S. (2007). A design science research methodology for information systems research. Journal of Management Information Systems, 24(3), 45\u201377. https://doi.org/10.2753/MIS0742-1222240302",
    "Richards, D., & Richardson, T. (2012). Computer-based psychological treatments for depression: A systematic review and meta-analysis. Clinical Psychology Review, 32(4), 329\u2013342. https://doi.org/10.1016/j.cpr.2012.02.004",
    "Spitzer, R. L., Kroenke, K., Williams, J. B. W., & L\u00f6we, B. (2006). A brief measure for assessing generalized anxiety disorder: The GAD-7. Archives of Internal Medicine, 166(10), 1092\u20131097. https://doi.org/10.1001/archinte.166.10.1092",
    "Torous, J., J\u00e4n Myrick, K., Rauseo-Ricupero, N., & Firth, J. (2020). Digital mental health and COVID-19: Using technology today to accelerate the curve on access and quality tomorrow. JMIR Mental Health, 7(3), e18848. https://doi.org/10.2196/18848",
    "United Republic of Tanzania (URT). (2022). Health sector report. Ministry of Health.",
    "World Health Organization. (2021). Mental health atlas 2020. World Health Organization.",
    "World Health Organization. (2022). World mental health report: Transforming mental health for all. World Health Organization.",
]
for r in refs:
    p = doc.add_paragraph(r)
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.5)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# ============================================================================= APPENDICES
page_break()
h1("APPENDICES")

# ----------------------------------------------------------------------------- Appendix A
h2("Appendix A: Questionnaire")
para(
    "The instruments below were used for requirements gathering, for the screening "
    "function of the system, and for usability evaluation. The screening instruments "
    "(PHQ-9, GAD-7 and SDQ) are standardized, validated tools and are reproduced here "
    "exactly as administered by the system. Registered users provide account details "
    "during sign-up; assessment demographic data (age group and sex) are collected "
    "separately before each screening."
)

para("Section A: Non-Identifying Background (collected before screening)", bold=True)
bullet("Age group: 12\u201315, 16\u201320, 21\u201325, 26\u201330, 31\u201340, or 40+")
bullet("Sex: Male or Female")
bullet("Preferred language: English or Kiswahili")

para("Section B: PHQ-9 Depression Screening Instrument", bold=True)
para(
    "Instruction shown to the respondent: \u201cOver the last 2 weeks, how often have "
    "you been bothered by any of the following problems?\u201d Each item is scored 0 = "
    "Not at all, 1 = Several days, 2 = More than half the days, 3 = Nearly every day.",
    italic=True,
)
simple_table(
    ["No.", "Item"],
    [
        ["1", "Little interest or pleasure in doing things"],
        ["2", "Feeling down, depressed, or hopeless"],
        ["3", "Trouble falling or staying asleep, or sleeping too much"],
        ["4", "Feeling tired or having little energy"],
        ["5", "Poor appetite or overeating"],
        ["6", "Feeling bad about yourself \u2014 or that you are a failure or have let yourself or your family down"],
        ["7", "Trouble concentrating on things, such as reading the newspaper or watching television"],
        ["8", "Moving or speaking so slowly that other people could have noticed; or the opposite \u2014 being so fidgety or restless that you have been moving around a lot more than usual"],
        ["9", "Thoughts that you would be better off dead or of hurting yourself in some way"],
    ],
)

para("Section C: GAD-7 Anxiety Screening Instrument", bold=True)
para(
    "Instruction and scoring are identical to the PHQ-9 (0\u20133 per item over the "
    "last two weeks).",
    italic=True,
)
simple_table(
    ["No.", "Item"],
    [
        ["1", "Feeling nervous, anxious, or on edge"],
        ["2", "Not being able to stop or control worrying"],
        ["3", "Worrying too much about different things"],
        ["4", "Trouble relaxing"],
        ["5", "Being so restless that it is hard to sit still"],
        ["6", "Becoming easily annoyed or irritable"],
        ["7", "Feeling afraid, as if something awful might happen"],
    ],
)

para("Section D: SDQ Children's Screening Instrument", bold=True)
para(
    "For children, the platform administers the Strengths and Difficulties "
    "Questionnaire (SDQ), the standard 25-item instrument completed by a parent or "
    "guardian on behalf of a child. Each item is rated \u201cNot True\u201d, "
    "\u201cSomewhat True\u201d or \u201cCertainly True\u201d, and the items are grouped "
    "into five subscales: emotional symptoms, conduct problems, hyperactivity/"
    "inattention, peer relationship problems, and prosocial behaviour. The instrument "
    "is used in full as published by Goodman (1997)."
)

para("Section E: Usability and Acceptance Questionnaire", bold=True)
para(
    "Administered after participants used the prototype. Each statement is rated on a "
    "five-point Likert scale (1 = Strongly Disagree to 5 = Strongly Agree).",
    italic=True,
)
simple_table(
    ["No.", "Statement", "1", "2", "3", "4", "5"],
    [
        ["1", "The system was easy to navigate and use.", "", "", "", "", ""],
        ["2", "The assessment questions were clear and easy to understand.", "", "", "", "", ""],
        ["3", "I received my results and feedback quickly.", "", "", "", "", ""],
        ["4", "The AI-generated feedback was helpful and supportive.", "", "", "", "", ""],
        ["5", "I felt my privacy and anonymity were respected.", "", "", "", "", ""],
        ["6", "The option to use Kiswahili was useful to me.", "", "", "", "", ""],
        ["7", "The conversational assistant (Mia) was helpful.", "", "", "", "", ""],
        ["8", "I would recommend the system to others.", "", "", "", "", ""],
    ],
)

# ----------------------------------------------------------------------------- Appendix B
page_break()
h2("Appendix B: Budget")
para(
    "The estimated budget for the development and evaluation of the system is "
    "summarized below (all amounts in Tanzanian Shillings, TZS)."
)
simple_table(
    ["No.", "Item", "Description", "Cost (TZS)"],
    [
        ["1", "Internet and data bundles", "Connectivity during development (6 months)", "180,000"],
        ["2", "Cloud hosting and domain", "Application server, database, and domain name", "250,000"],
        ["3", "Language model API usage", "DeepSeek / OpenAI API credits for feedback generation", "120,000"],
        ["4", "Equipment maintenance", "Laptop maintenance and depreciation", "150,000"],
        ["5", "Printing, binding and stationery", "Report production and documentation", "100,000"],
        ["6", "Contingency", "Miscellaneous and unforeseen expenses", "100,000"],
        ["", "Total", "", "900,000"],
    ],
)

# ----------------------------------------------------------------------------- Appendix C
h2("Appendix C: Work Plan")
para(
    "The project was carried out over six months. The work plan below shows the main "
    "activities and the months in which each was undertaken (\u25A0 indicates an active "
    "period)."
)
simple_table(
    ["Activity", "M1", "M2", "M3", "M4", "M5", "M6"],
    [
        ["Proposal and literature review", "\u25A0", "\u25A0", "", "", "", ""],
        ["Requirements analysis", "", "\u25A0", "\u25A0", "", "", ""],
        ["System and database design", "", "", "\u25A0", "", "", ""],
        ["Implementation (backend, frontend, AI)", "", "", "\u25A0", "\u25A0", "\u25A0", ""],
        ["Testing and evaluation", "", "", "", "", "\u25A0", ""],
        ["Report writing and submission", "", "", "", "", "\u25A0", "\u25A0"],
    ],
)

# ----------------------------------------------------------------------------- Appendix D
page_break()
h2("Appendix D: Program Code")
para(
    "Representative source-code listings from the implemented system are provided "
    "below. The complete source code is maintained in the project repository."
)

code_listing(
    "async submitToAI() {\n"
    "  this.loading = true;\n"
    "  try {\n"
    "    const res = await axios.post(`/api/assessment/phq9/`, {\n"
    "      responses: this.responses,\n"
    "      scores: this.responses.map(r => r.score),\n"
    "      lang_text: this.questions[0].text,\n"
    "      user_type: \"self\",\n"
    "      age_group: this.form.ageGroup,\n"
    "      sex: this.form.sex\n"
    "    });\n"
    "    const data = res.data;\n"
    "    this.aiMessageHTML = marked.parse(data.response || \"*No AI feedback.*\");\n"
    "    this.redirectLink = data.redirect_link || null;\n"
    "  } catch (error) {\n"
    "    this.aiMessageHTML = marked.parse(\"*Error contacting server.*\");\n"
    "  } finally {\n"
    "    this.loading = false;\n"
    "  }\n"
    "}",
    "Listing D.1: Frontend (Vue) \u2014 submitting PHQ-9 responses and rendering AI feedback.",
)

code_listing(
    "@api_view([\"POST\"])\n"
    "def phq9_assessment(request):\n"
    "    try:\n"
    "        scores = request.data.get(\"scores\", [])\n"
    "        responses = request.data.get(\"responses\", [])\n"
    "        lang_text = request.data.get(\"lang_text\")\n"
    "        user_type = request.data.get(\"user_type\")\n"
    "        age_group = request.data.get(\"age_group\")\n"
    "        sex = request.data.get(\"sex\")\n"
    "        questions = [r.get(\"question\") for r in responses]\n"
    "        if not scores or not questions or len(scores) != 9 or len(questions) != 9:\n"
    "            return Response({\"error\": \"Invalid or incomplete PHQ-9 input.\"},\n"
    "                            status=status.HTTP_400_BAD_REQUEST)\n"
    "        total_score = sum(scores)\n"
    "        lang = langdetect.detect(lang_text)\n"
    "        record = SelfAssessment.objects.create(\n"
    "            assessment_type='phq9', age_group=age_group, sex=sex, score=total_score)\n"
    "        prompt = generate_assessment_prompt(\n"
    "            user_type, total_score, 27, age_group, lang, zip(questions, scores))\n"
    "        client, model_name, temperature = get_model_client_for_feature(\n"
    "            \"phq9\", FeatureModelAssignment)\n"
    "        chat_response = client.chat.completions.create(\n"
    "            model=model_name,\n"
    "            messages=[{\"role\": \"user\", \"content\": prompt}],\n"
    "            temperature=temperature, stream=False)\n"
    "        result = chat_response.choices[0].message.content.strip()\n"
    "        redirect_link = f\"/followup/{record.id}\" if total_score >= 19 else None\n"
    "        return Response({\"score\": total_score, \"response\": result,\n"
    "                         \"redirect_link\": redirect_link}, status=status.HTTP_200_OK)\n"
    "    except Exception as e:\n"
    "        return Response({\"error\": str(e)},\n"
    "                        status=status.HTTP_500_INTERNAL_SERVER_ERROR)",
    "Listing D.2: Backend (Django REST) \u2014 PHQ-9 scoring and AI feedback endpoint.",
)

code_listing(
    "system_prompt = (\n"
    "    \"You are Mia, a supportive and empathetic AI mental health assistant \"\n"
    "    \"built for the MindCare platform. You specialize in psychological and \"\n"
    "    \"emotional wellbeing support.\\n\\n\"\n"
    "    \"Guidelines:\\n\"\n"
    "    \"1. Respond only in the user's detected language\\n\"\n"
    "    \"2. Use markdown formatting (**bold**, *italic*, [links](https://...))\\n\"\n"
    "    \"4. For critical situations, encourage the user to seek professional help\\n\"\n"
    "    \"5. Gently decline non-mental-health questions with a supportive tone\\n\"\n"
    "    \"7. Never offer medical diagnosis or treatment advice\"\n"
    ")",
    "Listing D.3: Backend \u2014 system prompt constraining the Mia assistant to a "
    "supportive, non-diagnostic role.",
)

# ----------------------------------------------------------------------------- Appendix E
page_break()
h2("Appendix E: Other Evidence")
para(
    "This appendix provides additional supporting evidence. The annotated screenshots "
    "of the running system are presented in Chapter Four (Figures 4.1\u20134.8). A "
    "representative example of the AI-generated feedback produced by the system for a "
    "moderate PHQ-9 score is reproduced below."
)
para("Sample AI-generated feedback (moderate depression score, 14 out of 27):", bold=True)
para(
    "\u201cThank you for completing the assessment. Your responses suggest a moderate "
    "level of depressive symptoms. This does not define you, and it is a helpful "
    "starting point for taking care of yourself: keep a gentle daily routine for "
    "sleep, meals and a little movement; reach out to one person you trust and share "
    "how you have been feeling; and notice small positive moments each day. If these "
    "feelings continue for more than two weeks or become heavier, please consider "
    "speaking with a counsellor or health worker. You are not alone, and support is "
    "available.\u201d",
    italic=True,
)
para(
    "Sample comparison systems referenced in Chapter Two (BetterHelp, Wysa, MindDoc, "
    "7 Cups and Headspace) are all publicly available online; screenshots of their "
    "interfaces may be inserted in Section 2.3 at the indicated positions."
)

# ----------------------------------------------------------------------------- save
import time

out_path = r"d:\Careen\frontend\Careen_Report_Updated.docx"
try:
    doc.save(out_path)
except PermissionError:
    out_path = r"d:\Careen\frontend\Careen_Report_Updated_%s.docx" % time.strftime("%H%M%S")
    doc.save(out_path)
    print("Original file was locked (open in Word); saved a fresh copy instead.")
print("Saved:", out_path)
print("Sections:", len(doc.sections))
print("Paragraphs:", len(doc.paragraphs))
